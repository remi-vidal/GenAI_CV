import re
import os
import PyPDF2 as pdf
import pdfplumber
import logging
import base64
from bson.binary import Binary
from docx import Document

def getResume(msg, cvs_folder):
    """
    Processes the attachments in the provided message to retrieve a resume.
    Args:
        msg: The message object containing attachments.
        cvs_folder: The folder where the resume should be saved.
    Raises:
        ValueError: If an attachment's name is None.
        Exception: If there are no attachments, or if an attachment's extension is not .pdf,
                   or if there is an error when getting the resume.
    Logs:
        Info: Logs the number of attachments if present.
        Error: Logs errors related to attachment name, unsupported extension,
               absence of attachments, and general errors when getting the resume.
    """

    try:
        if msg.attachments:
            logging.info("There is %d attachement(s)....", len(msg.attachments))

            for attachment in msg.attachments:
                if attachment.longFilename is None:
                    logging.error("Attachment name is None")
                    raise ValueError("Attachment name is None")

                _, extension = os.path.splitext(attachment.longFilename)

                if extension not in [".pdf", ".docx"]:
                    logging.error(f"Unsupported extension {extension}")
                    return None

                # Temporary save
                attachment.save(customFilename=attachment.longFilename)

                # Move the file to the correct folder
                saved_path = os.path.join(os.getcwd(), attachment.longFilename)
                final_path = os.path.join(cvs_folder, attachment.longFilename)

                # If the file already exists, remove it
                if os.path.exists(final_path):
                    logging.info(f"Replacing existing file: {final_path}")
                    os.remove(final_path)

                os.rename(saved_path, final_path)  # Move the file to the final folder
                return final_path  # Return the file path

        else:
            logging.error("No attachment found")
            return None

    except Exception as e:
        logging.error(f"Error when getting resume: {e}")
        return None

def extract_linkedin_infos(message):
    """
    Extracts LinkedIn information from a given message.
    This function searches for a specific pattern in the message body to extract
    the title and address associated with a LinkedIn profile.
    Args:
        message (extract_msg.msg_classes.message.Message): Parser for Microsoft Outlook message files
    Returns:
        tuple: A tuple containing the extracted title and address. If the pattern
               is not found, both title and address will be "N/A".
    """

    body = message.body

    # On split en utilisant \t\r\n comme séparateur
    parts = re.split(r'\t\r\n', body)

    # Vérification que la liste contient assez d'éléments
    if len(parts) > 4:
        title = parts[3].strip()
        address = parts[4].strip()
        
    else:
        title = address = "N/A"

    return title, address


def extract_text_from_pdf(file):
    """
    Extracts text from a PDF file. Using PyPDF2 if possible, otherwise using pdfplumber.

    Args:
        file (str): The path to the PDF file.

    Returns:
        str: The extracted text from the PDF, with newlines replaced by spaces and leading/trailing whitespace removed.
    """
    try : 
        reader = pdf.PdfReader(file)
        text = ""
        for _, page in enumerate(reader.pages):
            text += str(page.extract_text())
            print('on passe dans le try')

    except :
        with pdfplumber.open(file) as pdf_p:
            print('on passe dans lexcept')
            text = " ".join(page.extract_text() or "" for page in pdf_p.pages)
        

    text = text.replace("\x00", "").replace("\n", " ").strip()
    return text


def extract_text_from_docx(file):
    """
    Extracts text from a DOCX file, including text from tables and paragraphs.
    Args:
        file (str): The path to the DOCX file.
    Returns:
        str: The extracted text with table cells separated by tabs, paragraphs separated by spaces, 
             and non-breaking spaces replaced by regular spaces.
    """
    doc = Document(file)
    text = []

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text.append("\t".join(row_text))  # Better column separation

    # Extract text from paragraphs
    for para in doc.paragraphs:
        text.append(para.text)

    cleaned_text = (
        "\n".join(text)
        .replace("\xa0", " ")
        .replace("\n", " ")
        .replace("\t", " ")
        .strip()
        )
    return cleaned_text


def anonymize_cv(text_cv, noms_from_email):
    # 1. Extraction de l'email avant anonymisation
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text_cv)
    extracted_email = email_match.group(0) if email_match else None

    # 2. Extraction du numéro de téléphone avant anonymisation
    phone_pattern = re.compile(r'''
        (?<!\d)                     # Ne pas être précédé par un chiffre
        (?:\+?\d{1,3}[-.\s]?)?      # Code pays optionnel (ex: +33)
        (?!2\d{3}[-.\s])            # Ne pas commencer par 2 suivi de 3 chiffres
        (?:\(?[3-9]\d{1,3}\)?[-.\s]?)? # Code régional optionnel (ex: (01) ou 01)
        (?:\d{2,4}[-.\s]?){2,3}     # Groupes de 2 à 4 chiffres séparés par des tirets, points ou espaces
        \d{2,4}                     # Dernier groupe de 2 à 4 chiffres
        (?!\d)                      # Ne pas être suivi par un chiffre
    ''', re.VERBOSE)

    # Récupérer tous les matchs
    phone_matches = phone_pattern.findall(text_cv)

    # Filtrer les matchs pour éliminer ceux qui commencent par "1" ou "2" pour éliminer  les années
    filtered_matches = [match for match in phone_matches if not match.strip().startswith(('1', '2'))]
    extracted_phone =filtered_matches[0] if filtered_matches else None
    
    # 3. Suppression du nom et prénom, même s'ils sont collés à d'autres mots
    for nom in sorted(noms_from_email, key=len, reverse=True):  # Trier pour éviter les conflits
        text_cv = re.sub(rf'(?i){re.escape(nom)}', '[ANONYMISÉ]', text_cv)

    # 4. Suppression de l'email (remplacement après extraction)
    text_cv = re.sub(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', '[EMAIL]', text_cv)

    # 5. Suppression des numéros de téléphone
    if extracted_phone:
        text_cv = text_cv.replace(extracted_phone, '[TÉL]')
    
    # 6. Suppression de l'adresse postale (basique, peut être affiné)
    text_cv = re.sub(r'\d{1,5}\s+\w+(?:\s+\w+)*(?:,\s*\w+(?:\s+\w+)*)?,?\s*\d{5}', '[ADRESSE]', text_cv)

    return text_cv, extracted_email, extracted_phone  # On retourne le texte anonymisé + l'email et le numéro de téléphone extraits


def highlight_rows(row):
    color = ""
    if str(row["Diplôme"]).isdigit() and int(row["Diplôme"]) <= 2020:
        color = "background-color: lightgreen"
    if row["Freelance"] == "OUI":
        color = "background-color: lightblue"
    return [color] * len(row)


def guess_extension(cv_bytes):
    """ Détecte l'extension du fichier en fonction de sa signature binaire. """
    signatures = {
        b"%PDF": "pdf",  # PDF
        b"PK\x03\x04": "docx",  # DOCX (et autres fichiers ZIP)
        b"\xd0\xcf\x11\xe0": "doc"  # DOC (format binaire ancien)
    }
    for sig, ext in signatures.items():
        if cv_bytes.startswith(sig):
            return ext
    return "bin"  # Par défaut, inconnu

def generate_download_link(cv_binary, filename="cv"):
    """ Génère un lien de téléchargement HTML pour un fichier binaire. """
    if isinstance(cv_binary, Binary):  # Vérifier si c'est un objet Binary de MongoDB
        cv_bytes = cv_binary.decode()
    elif isinstance(cv_binary, bytes):
        cv_bytes = cv_binary
    else:
        return "✖️"

    file_extension = guess_extension(cv_bytes)
    full_filename = f"{filename}.{file_extension}"

    # Déterminer le type MIME
    mime_types = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword",
    }
    mime_type = mime_types.get(file_extension, "application/octet-stream")  # Fallback générique

    # Encodage base64
    b64 = base64.b64encode(cv_bytes).decode()
    href = f'<a href="data:{mime_type};base64,{b64}" download="{full_filename}">📄</a>'
    return href
